#!/usr/bin/env python3
"""Wypełnia szablon `wezwanie_WZ-2026-ANOW001_Nowak.docx` danymi WZ/2026/APS001 (Pielaszewski).

Wynik ma być identyczny z wezwaniem dla Adama Nowaka pod względem układu Worda (ten sam .docx),
tylko podmienione kwoty, tabele i sygnatura — zgodnie z `generate_wezwanie_pielaszewski_pdf.py`.
Wymaga: pip install python-docx
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Zainstaluj: pip install python-docx", file=sys.stderr)
    raise

SYGN = "WZ/2026/APS001"
TOTAL = "3820.73"
TERM_PAY = "30 kwietnia 2026"
DNI_PO = "3"

DUE_UNPAID = "25.07.2025"
DNI_UNPAID = 275
MAIN_350 = "350.00"
ODS_UNPAID = "41.53"
REK_1 = "170.92"
ODS_LATE = "10.80"
REK_19 = "3247.48"

ROWS_LATE: list[tuple[str, str, str, str, str, str, str]] = [
    ("8/02/2022", "249.00", "16.02.2022", "17.02.2022", "1 dzień", "0.11", "170.92"),
    ("18/09/2022", "249.00", "19.09.2022", "20.09.2022", "1 dzień", "0.11", "170.92"),
    ("27/08/2022", "249.00", "19.08.2022", "25.08.2022", "6 dni", "0.64", "170.92"),
    ("17/12/2022", "349.00", "13.12.2022", "14.12.2022", "1 dzień", "0.15", "170.92"),
    ("5/03/2023", "450.00", "13.03.2023", "16.03.2023", "3 dni", "0.58", "170.92"),
    ("19/06/2023", "350.00", "19.06.2023", "21.06.2023", "2 dni", "0.30", "170.92"),
    ("17/07/2023", "350.00", "17.07.2023", "19.07.2023", "2 dni", "0.30", "170.92"),
    ("15/08/2023", "350.00", "15.08.2023", "17.08.2023", "2 dni", "0.30", "170.92"),
    ("16/09/2023", "350.00", "12.09.2023", "18.09.2023", "6 dni", "0.91", "170.92"),
    ("8/11/2023", "350.00", "13.11.2023", "17.11.2023", "4 dni", "0.60", "170.92"),
    ("14/03/2024", "350.00", "13.03.2024", "18.03.2024", "5 dni", "0.76", "170.92"),
    ("5/05/2024", "350.00", "14.05.2024", "15.05.2024", "1 dzień", "0.15", "170.92"),
    ("7/08/2024", "350.00", "16.08.2024", "19.08.2024", "3 dni", "0.45", "170.92"),
    ("15/09/2024", "350.00", "13.09.2024", "16.09.2024", "3 dni", "0.45", "170.92"),
    ("20/10/2024", "350.00", "14.10.2024", "21.10.2024", "7 dni", "1.06", "170.92"),
    ("19/12/2024", "350.00", "22.12.2024", "27.12.2024", "5 dni", "0.76", "170.92"),
    ("16/02/2025", "350.00", "13.02.2025", "20.02.2025", "7 dni", "1.06", "170.92"),
    ("16/03/2025", "350.00", "13.03.2025", "24.03.2025", "11 dni", "1.66", "170.92"),
    ("20/05/2025", "350.00", "13.06.2025", "16.06.2025", "3 dni", "0.45", "170.92"),
]

SUMMARY_LEFT = [
    "Należność główna (1 faktura niezapłacona)",
    "Odsetki ustawowe za opóźnienie (15,75% p.a., 1 faktura niezapłacona)",
    "Rekompensata za koszty odzyskiwania (art. 10, 1 faktura niezapł., kurs NBP)",
    "Odsetki ustawowe za opóźnienie (16,00%/16,75%/15,75% p.a., 19 faktur spóźn.)",
    "Rekompensata za koszty odzyskiwania (art. 10, 19 faktur spóźn., kurs NBP)",
    "ŁĄCZNA KWOTA DO ZAPŁATY",
]
SUMMARY_RIGHT = [
    f"{MAIN_350} PLN",
    f"{ODS_UNPAID} PLN",
    f"{REK_1} PLN",
    f"{ODS_LATE} PLN",
    f"{REK_19} PLN",
    f"{TOTAL} PLN",
]


def _set_row_cells(table, row_idx: int, values: list[str]) -> None:
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        row.cells[i].text = val


def fill(doc_path: Path, out_docx: Path) -> None:
    d = Document(str(doc_path))

    d.paragraphs[2].text = f"Sygn.: {SYGN}"

    t0 = d.tables[0]
    _set_row_cells(
        t0,
        1,
        [
            "16/07/2025",
            f"{MAIN_350} PLN",
            DUE_UNPAID,
            f"{DNI_UNPAID} dni",
            f"{ODS_UNPAID} PLN",
            f"{REK_1} PLN",
            "Nieopłacona w całości",
        ],
    )

    t1 = d.tables[1]
    need = 1 + len(ROWS_LATE)
    while len(t1.rows) < need:
        t1.add_row()
    while len(t1.rows) > need:
        tr = t1.rows[-1]._tr
        t1._tbl.remove(tr)

    for i, x in enumerate(ROWS_LATE, start=1):
        _set_row_cells(
            t1,
            i,
            [x[0], f"{x[1]} PLN", x[2], x[3], x[4], f"{x[5]} PLN", f"{x[6]} PLN"],
        )

    t2 = d.tables[2]
    for r in range(6):
        t2.rows[r].cells[0].text = SUMMARY_LEFT[r]
        t2.rows[r].cells[1].text = SUMMARY_RIGHT[r]

    wyzn = (
        f"Wyznaczam termin zapłaty pełnej kwoty {TOTAL} PLN do dnia {TERM_PAY} "
        f"({DNI_PO} dni od daty niniejszego wezwania)."
    )
    d.paragraphs[19].text = wyzn

    t3 = d.tables[3]
    t3.rows[3].cells[1].text = SYGN
    t3.rows[4].cells[1].text = f"{TOTAL} PLN"

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out_docx))


def _soffice() -> str | None:
    for c in ("/Applications/LibreOffice.app/Contents/MacOS/soffice", "soffice"):
        if c.startswith("/"):
            p = Path(c)
            if p.is_file():
                return str(p)
        r = subprocess.run(["which", c], capture_output=True, text=True)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout.strip()
    return None


def try_pdf(in_docx: Path, out_pdf: Path) -> bool:
    soff = _soffice()
    if not soff:
        return False
    out_dir = in_docx.parent
    subprocess.run(
        [soff, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(in_docx)],
        check=False,
        capture_output=True,
        text=True,
    )
    expected = in_docx.with_suffix(".pdf")
    if expected.is_file() and expected.resolve() != out_pdf.resolve():
        shutil.move(str(expected), str(out_pdf))
    return out_pdf.is_file()


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    template = root / "templates" / "wezwanie_WZ-2026-ANOW001_Nowak.docx"
    if not template.is_file():
        print(f"Brak szablonu: {template}", file=sys.stderr)
        sys.exit(1)

    out_docx = root / "reports" / "wezwanie_WZ-2026-APS001_Pielaszewski.docx"
    fill(template, out_docx)
    print("Zapisano:", out_docx)

    out_pdf = root / "reports" / "wezwanie_WZ-2026-APS001_Pielaszewski.pdf"
    if try_pdf(out_docx, out_pdf):
        print("PDF (LibreOffice):", out_pdf)
    else:
        print(
            "PDF: otwórz plik .docx w Wordzie i Eksportuj do PDF, albo zainstaluj LibreOffice (soffice).",
            file=sys.stderr,
        )

    case = Path(
        "/Users/prometheus/Desktop/SPRAWY Z DŁUŻNIKAMI/wezwania do zapłaty wysłane/ADAM PIELASZEWSKI"
    )
    if case.is_dir():
        d_docx = case / "wezwanie_WZ-2026-APS001_Pielaszewski.docx"
        try:
            shutil.copy2(out_docx, d_docx)
            print("Skopiowano do sprawy:", d_docx)
        except OSError as e:
            print("Kopia do sprawy — pominięto:", e, file=sys.stderr)
        if out_pdf.is_file():
            d_pdf = case / "wezwanie_WZ-2026-APS001_Pielaszewski.pdf"
            try:
                shutil.copy2(out_pdf, d_pdf)
            except OSError as e:
                print("Kopia PDF do sprawy — pominięto:", e, file=sys.stderr)


if __name__ == "__main__":
    main()
